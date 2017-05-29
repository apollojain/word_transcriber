from docx import Document
import subprocess
import urllib
from flickrapi import FlickrAPI
import secret_stuff
import os.path

FLICKR_PUBLIC = secret_stuff.public
FLICKR_SECRET = secret_stuff.secret

def find_flickr_photo(keyword):
	flickr = FlickrAPI(FLICKR_PUBLIC, FLICKR_SECRET, format='parsed-json')
	extras='url_sq,url_t,url_s,url_q,url_m,url_n,url_z,url_c,url_l,url_o'
	cats = flickr.photos.search(text=keyword, per_page=5, extras=extras)
	photo_url = 'images/no-image-found.jpg'
	photos = cats['photos']['photo']
	if len(photos) > 0: 
		photo_url = photos[0]['url_c']
	else: 
		return photo_url
	i = 0
	savename = keyword
	while os.path.isfile('images/' + savename + str(i) + '.jpg'):
		i += 1
	result_path = 'images/' + savename + str(i) + '.jpg'
	urllib.urlretrieve(photo_url, result_path)
	return result_path


def process_string_to_doc(string, filename):
	document = Document()
	array = string.split('||')[1:]
	n = len(array)
	i = 0
	print array
	while i < n:
		if array[i] == 'small':
			document.add_heading(array[i + 1], level=1)
		elif array[i] == 'medium': 
			document.add_heading(array[i + 1], level=2)
		elif array[i] == 'large': 
			document.add_heading(array[i + 1], level=3) 
		elif array[i] == 'img':
			document.add_picture(array[i + 1])
		else: 
			document.add_paragraph(array[i + 1])
		i += 2
		document.save(filename)

def preview(string, filename):
	process_string_to_doc(string, filename)
	subprocess.call(('open', filename))

def save_url_from_web(address, filename):
	urllib.urlretrieve(address, filename)


if __name__ == '__main__':
	string = "||small|| My name is fred ||pr|| \n\tI live in a button factory ||img|| https://farm5.staticflickr.com/4198/34142658203_9f41601bdc_c.jpg"
	process_string_to_doc(string, 'sample.docx')
	# print find_flickr_photo('cat')