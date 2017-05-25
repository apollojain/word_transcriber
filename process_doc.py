from docx import Document

def process_string_to_doc(string, filename):
	document = Document()
	array = string.split('||')[1:]
	n = len(array)
	i = 0
	print array
	while i < n:
		if array[i] == 'small':
			document.add_heading(array[i + 1], level=1)
		elif array[i] == 'medium': 
			document.add_heading(array[i + 1], level=2)
		elif array[i] == 'large': 
			document.add_heading(array[i + 1], level=3) 
		elif array[i] == 'img':
			document.add_picture(array[i + 1], width=Inches(5.0))
		else: 
			document.add_paragraph(array[i + 1])
		i += 2
		document.save(filename)


if __name__ == '__main__':
	string = "||small|| My name is fred ||pr|| \n\tI live in a button factory"
	process_string_to_doc(string, 'sample.docx')